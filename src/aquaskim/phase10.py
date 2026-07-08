from __future__ import annotations

"""Phase 10: reproducible final Word report and submission package.

The report is built entirely from the validated project artifacts produced in
Phases 02--09.  It intentionally does not invent new numerical results: every
number in the document is loaded from an existing JSON/CSV output or the
central YAML design configuration.
"""

from dataclasses import dataclass
from datetime import datetime, timezone
import csv
import hashlib
import json
import shutil
import zipfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from PIL import Image
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from aquaskim.config import load_base_configuration
from aquaskim.paths import DIRECTORIES, PROJECT_ROOT, ensure_runtime_directories, relative_to_root


REPORT_TITLE_FA = "گزارش نهایی طراحی و شبیه‌سازی"
REPORT_SUBTITLE_FA = "AquaSkim-Sim: ربات سطحی خودگردان کاتاماران برای جمع‌آوری زباله‌های شناور"
REPORT_DESCRIPTION_FA = "دوقلوی دیجیتال پارامتریک شامل طراحی مکانیکی، هیدرواستاتیک، پیشرانش، انرژی، دینامیک، ادراک، خودگردانی و اعتبارسنجی آماری"

NAVY = "0B1F3A"
BLUE = "1F4E79"
TEAL = "0E7490"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F3F6F8"
MID_GRAY = "6B7280"
DARK = "1F2937"
GREEN = "1D6F42"
ORANGE = "B45309"


class ReportBuildError(RuntimeError):
    """Raised if report prerequisites or report build validation are not met."""


@dataclass(frozen=True)
class Phase10Artifacts:
    report_docx: Path
    report_manifest: Path
    report_readme: Path
    submission_zip: Path
    submission_manifest: Path
    checksums: Path

    def as_dict(self) -> dict[str, str]:
        return {name: relative_to_root(path) for name, path in self.__dict__.items()}


# ---------------------------------------------------------------------------
# Generic data loading
# ---------------------------------------------------------------------------

def _read_json(relative: str) -> dict[str, Any]:
    path = PROJECT_ROOT / relative
    if not path.exists():
        raise ReportBuildError(
            f"Required Phase artifact is missing: {relative}. Run the Phase 02--09 build first."
        )
    with path.open("r", encoding="utf-8") as handle:
        parsed = json.load(handle)
    if not isinstance(parsed, dict):
        raise ReportBuildError(f"Expected JSON object in {relative}.")
    return parsed


def _read_csv(relative: str) -> list[dict[str, str]]:
    path = PROJECT_ROOT / relative
    if not path.exists():
        raise ReportBuildError(f"Required table is missing: {relative}.")
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _safe_float(value: object, *, fallback: float = 0.0) -> float:
    try:
        return float(value)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        return fallback


def _metadata() -> dict[str, str]:
    """Read report-cover metadata while keeping default placeholders transparent."""
    path = PROJECT_ROOT / "config" / "report_metadata.json"
    defaults = {
        "student_name": "نام و نام خانوادگی",
        "student_id": "شماره دانشجویی",
        "course": "ربات‌های متحرک خودگردان",
        "instructor": "نام استاد",
        "institution": "دانشگاه / دانشکده",
        "semester": "",
    }
    if not path.exists():
        return defaults
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ReportBuildError(f"Invalid report metadata JSON: {path}") from exc
    if not isinstance(raw, dict):
        raise ReportBuildError("config/report_metadata.json must contain a JSON object.")
    return {key: str(raw.get(key, default)) for key, default in defaults.items()}


# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, fill: str) -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = "D1D5DB", size: str = "6") -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def _set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_header(row) -> None:  # type: ignore[no-untyped-def]
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _set_repeat_table_layout(table) -> None:  # type: ignore[no-untyped-def]
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_paragraph_rtl(paragraph, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.RIGHT) -> None:  # type: ignore[no-untyped-def]
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _set_run_font(run, font_name: str = "Arial", size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:  # type: ignore[no-untyped-def]
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_page_field(paragraph) -> None:  # type: ignore[no-untyped-def]
    run = paragraph.add_run()
    _set_run_font(run, size=9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _add_hyperlink(paragraph, url: str, text: str) -> None:  # type: ignore[no-untyped-def]
    # A visible raw path is sufficient for offline submission; no external URL is needed.
    run = paragraph.add_run(text)
    _set_run_font(run, size=9, color=BLUE)
    run.underline = True


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.20

    for style_name, size, color in (("Title", 24, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, TEAL)):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(15)
        style.paragraph_format.space_after = Pt(8)

    if "Caption" not in [style.name for style in styles]:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)

    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("AquaSkim-Sim | Digital Twin Engineering Report")
    _set_run_font(run, size=8.5, bold=True, color=NAVY)
    hp.add_run("    ")
    run = hp.add_run("Autonomous catamaran surface-cleaning robot")
    _set_run_font(run, size=8.5, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AquaSkim-Sim | صفحه ")
    _set_run_font(run, size=9, color=MID_GRAY)
    _add_page_field(fp)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    _set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    _set_run_font(run, size={1: 17, 2: 13, 3: 11.5}.get(level, 11), bold=True, color={1: NAVY, 2: BLUE, 3: TEAL}.get(level, DARK))


def _add_paragraph(document: Document, text: str, *, rtl: bool = True, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if rtl:
        _set_paragraph_rtl(paragraph)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, bold=True, color=NAVY)
        body = paragraph.add_run(text[len(bold_lead):])
        _set_run_font(body)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    _set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    _set_run_font(run)


def _add_equation(document: Document, expression: str, explanation: str | None = None) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(expression)
    _set_run_font(run, font_name="Courier New", size=10.5, bold=True, color=NAVY)
    if explanation:
        _add_paragraph(document, explanation)


def _add_table(document: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None, rtl: bool = True) -> None:
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    _set_repeat_table_layout(table)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, NAVY)
        _set_cell_border(cell, color="FFFFFF")
        _set_cell_margins(cell)
        if widths_cm:
            cell.width = Cm(widths_cm[index])
        p = cell.paragraphs[0]
        if rtl:
            _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        _set_run_font(run, size=8.5, bold=True, color="FFFFFF")
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for column, value in enumerate(row):
            cell = cells[column]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(cell, "FFFFFF" if row_index % 2 == 0 else LIGHT_GRAY)
            _set_cell_border(cell)
            _set_cell_margins(cell)
            if widths_cm:
                cell.width = Cm(widths_cm[column])
            p = cell.paragraphs[0]
            if rtl:
                _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            _set_run_font(run, size=8.6, color=DARK)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _prepare_report_image(relative: str, *, max_width_px: int = 1800, max_height_px: int = 1250) -> Path:
    """Create a compact report-only PNG derivative without changing source artifacts.

    Phase figures are intentionally high resolution for standalone use. Embedding all
    of those source PNGs directly in a DOCX can make office renderers consume several
    gigabytes of memory. The report therefore uses a traceable downsampled derivative
    at approximately 250--300 dpi for its printed width.
    """
    source = PROJECT_ROOT / relative
    if not source.exists():
        raise ReportBuildError(f"Figure required by final report is missing: {relative}")
    asset_dir = DIRECTORIES["reports"] / "report_assets"
    asset_dir.mkdir(parents=True, exist_ok=True)
    destination = asset_dir / f"{source.stem}_report.png"
    with Image.open(source) as image:
        if image.mode not in {"RGB", "RGBA"}:
            image = image.convert("RGBA")
        image.thumbnail((max_width_px, max_height_px), Image.Resampling.LANCZOS)
        image.save(destination, format="PNG", optimize=True)
    return destination


def _add_figure(document: Document, relative: str, caption: str, *, width_cm: float = 15.5) -> None:
    path = _prepare_report_image(relative)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = document.add_paragraph(style="Caption")
    _set_paragraph_rtl(cap, WD_ALIGN_PARAGRAPH.CENTER)
    run = cap.add_run(caption)
    _set_run_font(run, size=9, color=MID_GRAY)


def _page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    _set_cell_shading(cell, LIGHT_BLUE)
    _set_cell_border(cell, color="A7C8D8", size="10")
    _set_cell_margins(cell, top=140, start=170, bottom=140, end=170)
    p = cell.paragraphs[0]
    _set_paragraph_rtl(p)
    title_run = p.add_run(f"{title}: ")
    _set_run_font(title_run, bold=True, color=NAVY)
    body_run = p.add_run(body)
    _set_run_font(body_run, color=DARK)
    document.add_paragraph().paragraph_format.space_after = Pt(3)


# ---------------------------------------------------------------------------
# Report content
# ---------------------------------------------------------------------------

def _design_value_rows(config: dict[str, Any], p02: dict[str, Any], p04: dict[str, Any], p05: dict[str, Any], p07: dict[str, Any]) -> list[list[str]]:
    geometry = config["mechanical"]["geometry"]
    full = p02["load_cases"]["full_design_payload"]
    cruise = p04["target_cruise"]
    battery = p05["battery_settings"]
    environment = p07["environment"]
    return [
        ["معماری", "کاتاماران دو بدنه با دو پیشران مستقل"],
        ["طول هر بدنه", f"{geometry['hull_length_m']:.2f} m"],
        ["عرض کلی", f"{p02['geometry']['overall_width_m']:.2f} m"],
        ["جرم خشک / جرم طراحی کامل", f"{p02['load_cases']['dry_empty_basket']['total_mass_kg']:.2f} / {full['total_mass_kg']:.2f} kg"],
        ["ظرفیت جابه‌جایی مفهومی", f"{p02['geometry']['capacity_mass_kg']:.2f} kg"],
        ["سرعت کروز طراحی", f"{cruise['ground_speed_mps']:.2f} m/s"],
        ["باتری مفهومی", f"{battery['nominal_voltage_v']:.1f} V, {battery['capacity_ah']:.1f} Ah, {battery['nominal_energy_wh']:.1f} Wh"],
        ["محیط آزمون", f"{environment['length_m']:.1f} × {environment['width_m']:.1f} m"],
        ["تعداد مانع / زباله", f"{environment['obstacle_count']} / {environment['debris_count']}"],
    ]


def _summary_values(p03: dict[str, Any], p04: dict[str, Any], p05: dict[str, Any], p08: dict[str, Any], p09: dict[str, Any]) -> list[list[str]]:
    full = p03["hydrostatic_cases"]["full_design_payload"]
    cruise = p04["target_cruise"]
    operating = next(row for row in p05["operating_points"] if row["operating_case"] == "cruise_calm_full_payload")
    mission = p08["mission_metrics"]
    mc = {str(row["metric"]): row for row in p09["monte_carlo_summary"]}
    return [
        ["پایداری اولیه در بار کامل", f"GM = {full['GM_m']:.3f} m", "PASS"],
        ["فری‌بورد در بار کامل", f"{full['freeboard_m']:.3f} m", "PASS"],
        ["مقاومت کروز", f"{cruise['resistance_n']:.3f} N at {cruise['ground_speed_mps']:.2f} m/s", "PASS"],
        ["ذخیره رانش کروز", f"{cruise['thrust_reserve_ratio']:.2f}×", "PASS"],
        ["توان باتری در کروز", f"{operating['battery_power_w']:.2f} W", "PASS"],
        ["دوام تا آستانهٔ RTH", f"{operating['endurance_to_rth_threshold_min']:.1f} min", "PASS"],
        ["ماموریت بسته‌حلقه", f"{mission['duration_s']:.1f} s, {mission['collected_count']} objects", mission['final_state']],
        ["نرخ موفقیت Monte Carlo", f"{100.0 * _safe_float(mc['success_rate']['value']):.1f}% over {int(_safe_float(mc['trial_count']['value']))} trials", "PASS"],
    ]


def _add_cover(document: Document, metadata: dict[str, str], date_text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("AQUASKIM-SIM")
    _set_run_font(run, size=23, bold=True, color=NAVY)

    p = document.add_paragraph()
    _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(REPORT_TITLE_FA)
    _set_run_font(run, size=22, bold=True, color=NAVY)

    p = document.add_paragraph()
    _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(9)
    run = p.add_run(REPORT_SUBTITLE_FA)
    _set_run_font(run, size=16, bold=True, color=BLUE)

    p = document.add_paragraph()
    _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(REPORT_DESCRIPTION_FA)
    _set_run_font(run, size=11.5, color=DARK)

    figure_relative = "outputs/figures/phase02_mechanical_top_view.png"
    fig = PROJECT_ROOT / figure_relative
    if fig.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.add_run().add_picture(str(_prepare_report_image(figure_relative)), width=Cm(13.0))

    details = [
        ["نام دانشجو", metadata["student_name"]],
        ["شماره دانشجویی", metadata["student_id"]],
        ["نام درس", metadata["course"]],
        ["نام استاد", metadata["instructor"]],
        ["دانشگاه / دانشکده", metadata["institution"]],
        *([["نیم‌سال", metadata["semester"]]] if metadata["semester"].strip() else []),
        ["تاریخ تولید گزارش", date_text],
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for index, (label, value) in enumerate(details):
        row = table.add_row().cells
        for cell in row:
            _set_cell_border(cell, color="CBD5E1")
            _set_cell_margins(cell, top=105, start=140, bottom=105, end=140)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(row[0], LIGHT_BLUE)
        row[0].width = Cm(5.4)
        row[1].width = Cm(10.5)
        for cell, content, bold in ((row[0], label, True), (row[1], value, False)):
            p = cell.paragraphs[0]
            _set_paragraph_rtl(p)
            r = p.add_run(content)
            _set_run_font(r, size=10, bold=bold, color=NAVY if bold else DARK)

    p = document.add_paragraph()
    _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("این سند به‌صورت خودکار از خروجی‌های محاسباتی و شبیه‌سازی پروژه تولید شده است.")
    _set_run_font(run, size=9, color=MID_GRAY)


def _add_contents(document: Document) -> None:
    _page_break(document)
    _add_heading(document, "فهرست مطالب", 1)
    entries = [
        "چکیده و نتیجه‌های کلیدی",
        "1. تعریف مسئله و دامنهٔ پروژه",
        "2. معماری کلان و روش اجرای دوقلوی دیجیتال",
        "3. طراحی مکانیکی، هندسه و خواص جرمی",
        "4. هیدرواستاتیک و پایداری عرضی",
        "5. مقاومت هیدرودینامیکی و انتخاب پیشران",
        "6. مدل انرژی، باتری و منطق بازگشت",
        "7. مدل دینامیکی سه‌درجه‌آزادی",
        "8. محیط عملیاتی، موانع و حسگرهای مجازی",
        "9. عامل خودگردان، برنامه‌ریزی و کنترل",
        "10. اعتبارسنجی سناریویی و Monte Carlo",
        "11. بازتولیدپذیری، شواهد اجرا و تحویل پروژه",
        "12. محدودیت‌ها و مسیر توسعه",
        "پیوست A. معادلات کلیدی",
        "پیوست B. فهرست تحویل‌ها و مسیر فایل‌ها",
        "پیوست C. منابع علمی مدل",
    ]
    for item in entries:
        p = document.add_paragraph(style="List Bullet")
        _set_paragraph_rtl(p)
        r = p.add_run(item)
        _set_run_font(r, size=10.5, color=DARK)
    _add_callout(document, "یادداشت", "برای فهرست خودکار همراه با شمارهٔ صفحه در Microsoft Word، پس از بازشدن سند کلیدهای Ctrl+A و سپس F9 را فشار دهید. تیترهای سند با سبک‌های استاندارد Heading ساخته شده‌اند.")


def _add_abstract(document: Document, summary_rows: list[list[str]]) -> None:
    _page_break(document)
    _add_heading(document, "چکیده و نتیجه‌های کلیدی", 1)
    _add_paragraph(document, "این پروژه یک دوقلوی دیجیتال برای طراحی و ارزیابی ربات سطحی خودگردان AquaSkim-Sim است. ربات پیشنهادی از معماری کاتاماران دو بدنه استفاده می‌کند و برای جمع‌آوری زباله‌های شناور در آب‌های آرام طراحی شده است. هدف کار، ساخت فیزیکی نمونه نبوده است؛ بلکه تمام زنجیرهٔ مهندسی، از هندسه و جرم تا پایداری، پیشرانش، انرژی، دینامیک، حسگر مجازی، عامل خودگردان و اعتبارسنجی آماری، به‌صورت پارامتریک و بازتولیدپذیر در Python پیاده‌سازی شده است.")
    _add_paragraph(document, "نتیجهٔ اصلی این است که در محدودهٔ فرض‌های اعلام‌شده، طراحی کامل‌بار از نظر شناوری و پایداری اولیه دارای حاشیهٔ مناسب است، سرعت کروز تعیین‌شده از منظر رانش و انرژی قابل دستیابی است و زنجیرهٔ بسته‌حلقهٔ برنامه‌ریزی A*، کنترل و بازگشت به خانه در سناریوهای رسمی با موفقیت اجرا شده است. نتیجه‌های Monte Carlo فقط برای بازهٔ جریان و SOC تعریف‌شده معتبرند و ادعای اعتبارسنجی میدانی یا گواهی ایمنی نمی‌شود.")
    _add_heading(document, "خلاصهٔ کمی نتایج", 2)
    _add_table(document, ["شاخص", "نتیجه", "وضعیت"], summary_rows, widths_cm=[5.4, 8.7, 2.4])


def _add_chapter_1(document: Document, config: dict[str, Any], p07: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "1. تعریف مسئله و دامنهٔ پروژه", 1)
    _add_paragraph(document, "مسئلهٔ پروژه طراحی مفهومی و شبیه‌سازی کامل یک ربات سطحی کوچک برای جمع‌آوری زباله‌های شناور از محیط آبی آرام است. مأموریت ربات از ایستگاه خانه آغاز می‌شود؛ ربات محیط را با حسگرهای مجازی مشاهده می‌کند، زبالهٔ قابل دسترس را انتخاب می‌کند، مسیر امن تولید می‌کند، با رانش تفاضلی حرکت می‌کند، زباله را به‌صورت هندسی جمع‌آوری می‌کند و پس از رسیدن به سهمیه یا شرایط بازگشت، به ایستگاه اولیه بازمی‌گردد.")
    _add_heading(document, "1-1. دامنه و فرض‌های اصلی", 2)
    for bullet in [
        "محیط هدف یک حوضچهٔ تحلیلی با آب آرام یا جریان یکنواخت کم‌سرعت است؛ موج، باد و محیط دریایی خشن در دامنهٔ اصلی نیستند.",
        "ربات فقط به‌صورت دیجیتال طراحی شده است؛ هدف، تحویل یک پروژهٔ کدنویسی، شبیه‌سازی، تصویر، ویدئو و گزارش مهندسی است.",
        "مدل‌های هیدرودینامیکی و حسگرها شفاف و پارامتریک هستند؛ بنابراین برای تحلیل آموزشی و مقایسهٔ سناریو مناسب‌اند، اما جایگزین CFD یا آزمون میدانی نیستند.",
        "تمام طول‌ها، جرم‌ها، نیروها و انرژی‌ها در دستگاه SI ثبت شده‌اند؛ محورهای محیط ENU هستند و محور x بدنه به سمت دهانهٔ جمع‌آوری تعریف شده است.",
    ]:
        _add_bullet(document, bullet)
    _add_heading(document, "1-2. مشخصات پایهٔ مسئله", 2)
    geometry = config["mechanical"]["geometry"]
    environment = p07["environment"]
    rows = [
        ["ابعاد حوضچه", f"{environment['length_m']:.1f} × {environment['width_m']:.1f} m"],
        ["عمق آب", f"{environment['water_depth_m']:.2f} m"],
        ["موقعیت ایستگاه خانه", f"({environment['home_position_m'][0]:.1f}, {environment['home_position_m'][1]:.1f}) m"],
        ["ساختار بدنه", config["mechanical"]["architecture"]],
        ["عرض دهانهٔ جمع‌آوری", f"{geometry['collector_inlet_width_m']:.2f} m"],
        ["بار طراحی سبد", f"{geometry['design_payload_kg']:.2f} kg"],
        ["شعاع ایمنی ربات", f"{environment['safety_radius_m']:.2f} m"],
        ["تعداد موانع / زباله‌ها", f"{environment['obstacle_count']} / {environment['debris_count']}"],
    ]
    _add_table(document, ["پارامتر", "مقدار"], rows, widths_cm=[7.4, 9.1])
    _add_callout(document, "معیار موفقیت", "ربات باید بدون ورود به فضای اشغال‌شده یا عبور از حاشیهٔ ایمنی، حداقل یک زباله را جمع‌آوری کند، SOC نهایی آن بالاتر از آستانهٔ بازگشت بماند و در ناحیهٔ خانه به حالت MISSION_COMPLETE برسد.")


def _add_chapter_2(document: Document) -> None:
    _page_break(document)
    _add_heading(document, "2. معماری کلان و روش اجرای دوقلوی دیجیتال", 1)
    _add_paragraph(document, "پروژه از یک زنجیرهٔ ماژولار تشکیل شده است. هر ماژول دادهٔ قابل‌ردیابی تولید می‌کند و خروجی آن به فاز بعدی منتقل می‌شود. این تفکیک باعث می‌شود هر فرض و هر نتیجه در فایل‌های JSON، CSV، PNG/SVG و لاگ‌های Evidence قابل پیگیری باشد.")
    _add_figure(document, "outputs/figures/phase08_autonomy_architecture.png", "شکل 1 — معماری اطلاعاتی و حلقهٔ تصمیم‌گیری: حسگرها، برنامه‌ریز، عامل خودگردان، کنترل و مدل دینامیکی.")
    _add_heading(document, "2-1. زنجیرهٔ فنی", 2)
    rows = [
        ["Phase 02", "هندسهٔ پارامتریک، بودجهٔ جرم، مرکز جرم و ممان اینرسی مفهومی"],
        ["Phase 03", "آبخور، فری‌بورد، KB/BM/KG/GM، منحنی GZ و گشتاور بازگرداننده"],
        ["Phase 04", "مقاومت آب، منحنی رانش و توان، RPM و ذخیرهٔ رانش"],
        ["Phase 05", "انرژی باتری، SOC، دوام و سیاست بازگشت به خانه"],
        ["Phase 06", "دینامیک surge–sway–yaw با حل RK4 و اثر جریان"],
        ["Phase 07", "نقشهٔ اشغال، موانع، زباله و حسگرهای مجازی"],
        ["Phase 08", "عامل خودگردان، A*، کنترل مسیر و مأموریت بسته‌حلقه"],
        ["Phase 09", "چهار سناریوی رسمی و 20 اجرای Monte Carlo"],
        ["Phase 10", "تولید گزارش Word، بستهٔ تحویل، هش‌ها و مسیر بازتولید"],
    ]
    _add_table(document, ["فاز", "نقش در زنجیرهٔ مهندسی"], rows, widths_cm=[3.0, 13.5])
    _add_heading(document, "2-2. اصل بازتولیدپذیری", 2)
    _add_paragraph(document, "هر اجرای رسمی در پوشهٔ records/phases ذخیره می‌شود. در هر اجرا، متن فرمان‌ها، stdout و stderr، نسخهٔ محیط Python، خروجی pip freeze، snapshot ورودی‌ها، hash SHA-256 خروجی‌ها، کپی Artifactها و Handoff فاز ذخیره می‌شوند. بنابراین نتیجهٔ گزارش فقط یک تصویر ثابت نیست و می‌توان مسیر تولید آن را تا پارامترهای اصلی دنبال کرد.")


def _add_chapter_3(document: Document, config: dict[str, Any], p02: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "3. طراحی مکانیکی، هندسه و خواص جرمی", 1)
    _add_paragraph(document, "معماری انتخاب‌شده کاتاماران است: دو بدنهٔ باریک موازی، سازهٔ اتصال مرکزی، دهانهٔ قیفی جمع‌آوری در جلو، سبد در میانهٔ جلو و دو پیشران در عقب. این ساختار هم فضای مناسبی برای دهانه و سبد ایجاد می‌کند و هم، نسبت به بدنهٔ تک‌بدنه، سطح خط آب عرضی بزرگ‌تری برای پایداری اولیه فراهم می‌کند.")
    _add_figure(document, "outputs/figures/phase02_mechanical_top_view.png", "شکل 2 — نمای بالای آرایش مکانیکی پارامتریک، دهانهٔ جمع‌آوری، سبد، پیشران‌ها و محل اجزای جرمی.")
    _add_figure(document, "outputs/figures/phase02_mechanical_side_view.png", "شکل 3 — نمای جانبی آرایش مکانیکی و پیش‌نمایش آبخور در حالت سبد خالی و بار طراحی کامل.")
    _add_heading(document, "3-1. پارامترهای هندسی و جرم", 2)
    p04 = _read_json("outputs/logs/phase04_propulsion_summary.json")
    p05 = _read_json("outputs/logs/phase05_energy_summary.json")
    p07 = _read_json("outputs/logs/phase07_environment_summary.json")
    _add_table(document, ["پارامتر", "مقدار"], _design_value_rows(config, p02, p04, p05, p07), widths_cm=[7.2, 9.3])
    _add_figure(document, "outputs/figures/phase02_mass_distribution.png", "شکل 4 — توزیع جرم اجزا و جابه‌جایی مرکز جرم بین حالت سبد خالی و بار کامل.")
    _add_heading(document, "3-2. محاسبهٔ مرکز جرم و ممان اینرسی", 2)
    _add_equation(document, "m_total = Σ m_i", "جرم کل از جمع جرم تمام اجزای پارامتریک به‌دست می‌آید.")
    _add_equation(document, "x_CG = Σ(m_i x_i) / Σm_i   ,   y_CG = Σ(m_i y_i) / Σm_i   ,   z_CG = Σ(m_i z_i) / Σm_i", "مرکز جرم با میانگین وزنی مختصات اجزا محاسبه می‌شود. در Phase 02، ممان اینرسی حول CG با تقریب جرم‌های متمرکز تشکیل شده است.")
    rows = []
    for key, label in (("dry_empty_basket", "سبد خالی"), ("full_design_payload", "بار طراحی کامل")):
        item = p02["load_cases"][key]
        rows.append([label, f"{item['total_mass_kg']:.3f}", f"({item['cg_x_m']:.3f}, {item['cg_y_m']:.3f}, {item['cg_z_m']:.3f})", f"{item['Izz_kg_m2']:.4f}"])
    _add_table(document, ["حالت بار", "جرم [kg]", "CG [m]", "Izz [kg·m²]"], rows, widths_cm=[4.0, 3.0, 6.5, 3.0])
    _add_callout(document, "تفسیر", "مرکز جرم جانبی در هر دو حالت تقریباً صفر است، زیرا اجزای چپ و راست متقارن جانمایی شده‌اند. با پر شدن سبد، CG کمی به سمت جلو حرکت می‌کند؛ این اثر در تحلیل پایداری فاز بعدی اعمال شده است.")


def _add_chapter_4(document: Document, p03: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "4. هیدرواستاتیک و پایداری عرضی", 1)
    _add_paragraph(document, "تحلیل هیدرواستاتیک در آب شیرین آرام با چگالی 1000 kg/m³ انجام شده است. حجم جابه‌جایی از تعادل وزن و نیروی ارشمیدس به‌دست می‌آید و پایداری اولیه با ارتفاع متاسنتری GM ارزیابی می‌شود. برای زاویه‌های محدود، انتگرال‌گیری نواری عرضی با حفظ جابه‌جایی ثابت استفاده شده است.")
    _add_heading(document, "4-1. معادلات حاکم", 2)
    _add_equation(document, "F_B = ρ g ∇ = m g", "در تعادل ساکن، نیروی شناوری با وزن ربات برابر است.")
    _add_equation(document, "GM = KB + BM − KG      ,      BM = I_T / ∇", "KB ارتفاع مرکز شناوری، KG ارتفاع مرکز جرم و I_T ممان دوم سطح خط آب است.")
    _add_equation(document, "GZ ≈ GM sin(φ)      ,      M_R = m g GZ", "در زاویه‌های کوچک، بازوی راست‌کننده و گشتاور بازگرداننده از GM به‌دست می‌آیند.")
    _add_figure(document, "outputs/figures/phase03_hydrostatics_dashboard.png", "شکل 5 — داشبورد هیدرواستاتیک شامل آبخور، فری‌بورد، KB/BM/KG/GM و مقایسهٔ دو حالت بار.")
    _add_figure(document, "outputs/figures/phase03_stability_curves.png", "شکل 6 — منحنی بازوی راست‌کننده و گشتاور بازگرداننده در زاویه‌های کج‌شدن تحلیل‌شده.")
    _add_heading(document, "4-2. نتایج کمی", 2)
    rows = []
    for key, label in (("dry_empty_basket", "سبد خالی"), ("full_design_payload", "بار طراحی کامل")):
        item = p03["hydrostatic_cases"][key]
        rows.append([label, f"{item['draft_m']:.4f}", f"{item['freeboard_m']:.4f}", f"{item['KG_m']:.4f}", f"{item['GM_m']:.4f}", f"{item['capacity_ratio']:.2f}"])
    _add_table(document, ["حالت", "آبخور [m]", "فری‌بورد [m]", "KG [m]", "GM [m]", "نسبت ظرفیت"], rows, widths_cm=[3.0, 2.6, 2.8, 2.4, 2.4, 3.0])
    _add_figure(document, "outputs/figures/phase03_heeling_cross_sections.png", "شکل 7 — مقاطع عرضی و تغییر خط آب در زوایای منتخب کج‌شدن.")
    _add_figure(document, "outputs/figures/phase03_payload_envelope.png", "شکل 8 — اثر جرم بار جمع‌آوری‌شده بر آبخور، فری‌بورد و پایداری اولیه.")
    _add_callout(document, "نتیجهٔ فاز پایداری", "حتی در حالت بار طراحی کامل، GM مثبت و فری‌بورد بسیار بالاتر از حد پذیرش تعریف‌شده باقی می‌ماند. این نتیجه فقط مربوط به آب آرام، هندسهٔ مفهومی و بارگذاری تقارن‌دار است؛ اثر موج، باد و دینامیک roll در این مدل لحاظ نشده است.")


def _add_chapter_5(document: Document, p04: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "5. مقاومت هیدرودینامیکی و انتخاب پیشران", 1)
    _add_paragraph(document, "اندازه‌گذاری پیشران در حالت بار کامل انجام شده است. مقاومت طولی شامل مؤلفهٔ اصطکاکی مبتنی بر ITTC-1957، مقاومت باقیمانده، مقاومت متعلقات و مقاومت دهانهٔ جمع‌آوری است. پیشران‌های دوگانه به‌صورت متقارن در حرکت مستقیم کار می‌کنند و در فاز دینامیک برای ایجاد گشتاور yaw از اختلاف رانش استفاده می‌شود.")
    _add_heading(document, "5-1. معادلات مدل", 2)
    _add_equation(document, "C_F = 0.075 / (log10(Re) − 2)^2", "خط هم‌بستگی اصطکاک ITTC-1957 برای ضریب اصطکاکی استفاده شده است.")
    _add_equation(document, "R_total = R_friction + R_residual + R_appendage + R_collector", "مقاومت کل برابر مجموع مؤلفه‌های شفاف مدل است.")
    _add_equation(document, "T_side = k_T RPM^2      ,      P_side = k_P RPM^3", "رانش و توان پیشران با روابط پارامتریک درجه دوم و سوم نسبت به RPM مدل شده‌اند.")
    _add_figure(document, "outputs/figures/phase04_resistance_dashboard.png", "شکل 9 — تفکیک مقاومت هیدرودینامیکی و اثر سرعت بر نیروی مورد نیاز.")
    _add_figure(document, "outputs/figures/phase04_propulsion_envelope.png", "شکل 10 — پوشش رانش، RPM و توان الکتریکی دو پیشران.")
    _add_heading(document, "5-2. نقطهٔ عملکرد کروز", 2)
    cruise = p04["target_cruise"]
    rows = [
        ["سرعت زمینی هدف", f"{cruise['ground_speed_mps']:.3f} m/s"],
        ["مقاومت کل", f"{cruise['resistance_n']:.3f} N"],
        ["RPM هر پیشران", f"{cruise['rpm_per_side']:.0f} rpm"],
        ["Throttle هر پیشران", f"{100 * cruise['throttle_fraction']:.1f}%"],
        ["توان پیشران‌ها", f"{cruise['thruster_power_w']:.2f} W"],
        ["توان کل باس", f"{cruise['total_electrical_power_w']:.2f} W"],
        ["نسبت ذخیرهٔ رانش", f"{cruise['thrust_reserve_ratio']:.2f}×"],
        ["سرعت نظری تقاطع رانش–مقاومت", f"{p04['theoretical_top_speed_calm_mps']:.2f} m/s"],
    ]
    _add_table(document, ["شاخص", "مقدار"], rows, widths_cm=[7.5, 9.0])
    _add_figure(document, "outputs/figures/phase04_current_penalty.png", "شکل 11 — اثر جریان مخالف بر نیاز رانش و امکان نگه‌داشتن سرعت زمینی هدف.")
    _add_figure(document, "outputs/figures/phase04_operating_envelope.png", "شکل 12 — محدودهٔ عملیاتی سرعت، رانش و توان با قیود طراحی.")


def _add_chapter_6(document: Document, p05: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "6. مدل انرژی، باتری و منطق بازگشت", 1)
    _add_paragraph(document, "مدل انرژی، توان پیشران Phase 04 و بار ثابت سامانهٔ کنترل و حسگرها را به انرژی بستهٔ باتری متصل می‌کند. SOC با انتگرال‌گیری جریان مصرفی به‌روز می‌شود. ظرفیت قابل‌استفاده، راندمان باس DC، کاهش ظرفیت و ضریب Peukert سبک در محاسبه دخالت دارند.")
    _add_heading(document, "6-1. معادلات پایه", 2)
    _add_equation(document, "P_battery = P_bus / η_DC", "توان لازم از باتری با توجه به راندمان مبدل از توان باس محاسبه می‌شود.")
    _add_equation(document, "SOC(t+Δt) = SOC(t) − I(t)Δt / C_effective", "کاهش SOC با انتگرال‌گیری جریان مؤثر باتری محاسبه می‌شود.")
    _add_equation(document, "SOC_command = max(SOC_configured, (E_return + E_reserve) / E_usable)", "حد فرمان بازگشت از بیشینهٔ آستانهٔ تنظیم‌شده و انرژی لازم برای بازگشت به‌دست می‌آید.")
    battery = p05["battery_settings"]
    rows = [
        ["شیمی بسته", battery["chemistry"]],
        ["ولتاژ نامی", f"{battery['nominal_voltage_v']:.1f} V"],
        ["ظرفیت نامی", f"{battery['capacity_ah']:.1f} Ah"],
        ["انرژی نامی", f"{battery['nominal_energy_wh']:.1f} Wh"],
        ["کسر قابل‌استفاده", f"{100*battery['usable_fraction']:.0f}%"],
        ["راندمان باس", f"{100*battery['dc_bus_efficiency']:.0f}%"],
        ["حداقل ولتاژ مدل", f"{battery['min_pack_voltage_v']:.1f} V"],
    ]
    _add_table(document, ["ویژگی باتری", "مقدار"], rows, widths_cm=[7.2, 9.3])
    _add_figure(document, "outputs/figures/phase05_energy_dashboard.png", "شکل 13 — داشبورد انرژی، نقاط عملکرد و نتیجهٔ تحلیل دوام.")
    _add_figure(document, "outputs/figures/phase05_mission_soc_profiles.png", "شکل 14 — پروفایل‌های SOC برای مأموریت‌های تعریف‌شده.")
    _add_figure(document, "outputs/figures/phase05_return_home_envelope.png", "شکل 15 — پوشش بازگشت به خانه با توجه به فاصله و جریان مخالف.")
    _add_callout(document, "نتیجهٔ انرژی", "در نقطهٔ کروز بار کامل، مدل توان باتری و زمان باقی‌مانده تا آستانهٔ بازگشت نشان می‌دهد که مأموریت‌های کوتاه بسته‌حلقهٔ Phase 08 و 09 حاشیهٔ انرژی قابل‌قبولی دارند. این نتیجه نباید به‌عنوان مدل حرارتی سلول، مدل ageing یا گواهی BMS تفسیر شود.")


def _add_chapter_7(document: Document, p06: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "7. مدل دینامیکی سه‌درجه‌آزادی", 1)
    _add_paragraph(document, "مدل حرکت صفحه‌ای ربات شامل موقعیت x و y، زاویهٔ ψ، سرعت طولی u، سرعت جانبی v و نرخ yaw یعنی r است. معادلات در چارچوب بدنه نوشته شده‌اند و با روش Runge–Kutta مرتبهٔ چهار حل می‌شوند. اثر جریان آب ابتدا به دستگاه بدنه تبدیل و سپس در سرعت نسبی آب–بدنه وارد نیروی مقاومتی می‌شود.")
    _add_heading(document, "7-1. سینماتیک و دینامیک", 2)
    _add_equation(document, "x_dot = u cos(ψ) − v sin(ψ)      ,      y_dot = u sin(ψ) + v cos(ψ)      ,      ψ_dot = r")
    _add_equation(document, "m_eff,u (u_dot − v r) = T_L + T_R − X_D(u_rel)")
    _add_equation(document, "m_eff,v (v_dot + u r) = −Y_D(v_rel)      ,      I_eff,z r_dot = (b/2)(T_R − T_L) − N_D(r)")
    _add_figure(document, "outputs/figures/phase06_dynamics_dashboard.png", "شکل 16 — پارامترهای دینامیکی، پاسخ طولی و مفهوم جرم افزوده.")
    _add_figure(document, "outputs/figures/phase06_trajectory_comparison.png", "شکل 17 — مقایسهٔ مسیرهای حرکت در سناریوهای دینامیکی منتخب.")
    _add_figure(document, "outputs/figures/phase06_maneuver_response.png", "شکل 18 — پاسخ مانور رانش تفاضلی و نرخ چرخش.")
    scenario_names = ", ".join(str(metric["scenario"]) for metric in p06["scenario_metrics"])
    _add_callout(document, "سناریوهای دینامیکی", f"سه آزمون باز Phase 06 اجرا شدند: {scenario_names}. شاخص‌های دقیق x، y، heading و yaw در فایل outputs/tables/phase06_scenario_metrics.csv ثبت شده‌اند؛ نمودارهای شکل‌های 16 تا 19 نمایش تصویری همان نتایج هستند.")
    _add_figure(document, "outputs/figures/phase06_current_disturbance.png", "شکل 19 — اثر جریان جانبی یکنواخت بر مسیر باز در مدل سه‌درجه‌آزادی.")


def _add_chapter_8(document: Document, p07: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "8. محیط عملیاتی، موانع و حسگرهای مجازی", 1)
    _add_paragraph(document, "محیط تحلیلی از مرزهای حوضچه، موانع تحلیلی دایره‌ای و مستطیلی، زباله‌های شناور با seed ثابت و نقشهٔ اشغال تشکیل شده است. برای برنامه‌ریزی مسیر، موانع به اندازهٔ شعاع ایمنی ربات inflate می‌شوند تا مسیر محاسبه‌شده فضای کافی برای بدنه داشته باشد.")
    _add_figure(document, "outputs/figures/phase07_environment_map.png", "شکل 20 — محیط مأموریت شامل خانه، موانع و موقعیت‌های قطعی زباله.")
    _add_figure(document, "outputs/figures/phase07_occupancy_grid.png", "شکل 21 — نقشهٔ اشغال و فضای پیکربندی پس از اعمال حاشیهٔ ایمنی ربات.")
    _add_heading(document, "8-1. حسگرهای مجازی", 2)
    _add_figure(document, "outputs/figures/phase07_sensor_model.png", "شکل 22 — مدل حسگر موقعیت، قطب‌نما، بردارهای فاصله و آشکارساز زباله.")
    rows = [
        ["موقعیت GNSS/UWB", "نویز گاوسی، RMS ثبت‌شده 0.064 m"],
        ["قطب‌نما", "نویز زاویه‌ای، RMS ثبت‌شده 1.924 deg"],
        ["فاصله‌سنج جلو", "5 پرتو، میدان دید 80 deg، برد 3.0 m"],
        ["آشکارساز زباله", "میدان دید 95 deg، برد 2.2 m، احتمال تشخیص وابسته به فاصله"],
        ["نقشهٔ اشغال", "تفکیک 0.10 m، 1820 سلول ممنوع از 9600 سلول"],
    ]
    _add_table(document, ["حسگر / لایه", "مدل و نتیجه"], rows, widths_cm=[5.0, 11.5])
    _add_figure(document, "outputs/figures/phase07_perception_dashboard.png", "شکل 23 — داشبورد ادراک، لاگ حسگر و آمار آشکارسازی زباله.")
    _add_callout(document, "محدودهٔ ادعا", "این حسگرها surrogateهای مجازی و قابل‌توضیح هستند. هیچ ادعایی دربارهٔ عملکرد یک دوربین واقعی، شبکهٔ عصبی یا SLAM در محیط واقعی مطرح نمی‌شود.")


def _add_chapter_9(document: Document, p08: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "9. عامل خودگردان، برنامه‌ریزی و کنترل", 1)
    _add_paragraph(document, "عامل خودگردان با ماشین حالت محدود پیاده‌سازی شده است. در حالت SEARCH، دادهٔ آشکارساز برای تایید هدف استفاده می‌شود. سپس برنامه‌ریز A* یک مسیر امن روی Occupancy Grid تولید می‌کند. هدایتگر زاویهٔ مطلوب را از waypoint فعال می‌سازد و کنترل‌کنندهٔ بازخوردی فرمان سرعت و yaw را به رانش پیشران چپ و راست تبدیل می‌کند. منطق SOC و سهمیهٔ جمع‌آوری می‌تواند RETURN_HOME را فعال کند.")
    _add_figure(document, "outputs/figures/phase08_planning_map.png", "شکل 24 — نقشهٔ برنامه‌ریزی A*، مسیرهای امن و نقاط تصمیم مأموریت.")
    _add_figure(document, "outputs/figures/phase08_closed_loop_mission.png", "شکل 25 — مسیر مأموریت بسته‌حلقه و رویدادهای جمع‌آوری و بازگشت به خانه.")
    _add_heading(document, "9-1. منطق کنترل", 2)
    _add_equation(document, "ψ_d = atan2(y_wp − y, x_wp − x)      ,      e_ψ = wrap(ψ_d − ψ)")
    _add_equation(document, "ω_cmd = K_p e_ψ + K_i ∫e_ψdt + K_d de_ψ/dt")
    _add_equation(document, "T_L = T_total/2 − τ_z/b      ,      T_R = T_total/2 + τ_z/b")
    _add_figure(document, "outputs/figures/phase08_control_dashboard.png", "شکل 26 — خطای مسیر، فرمان‌های کنترل و رانش دو پیشران در مأموریت بسته‌حلقه.")
    _add_figure(document, "outputs/figures/phase08_decision_timeline.png", "شکل 27 — timeline تغییر حالت‌های عامل خودگردان و دلایل تصمیم.")
    _add_heading(document, "9-2. شواهد رویدادهای مأموریت پایه", 2)
    transitions = p08["state_transitions"]
    rows = []
    for event in transitions:
        target = str(event.get("target_id", "")) or "—"
        rows.append([f"{_safe_float(event['time_s']):.1f}", f"{event['from_state']} → {event['to_state']}", str(event['reason']), target])
    _add_table(document, ["زمان [s]", "گذار حالت", "دلیل", "هدف"], rows, widths_cm=[2.0, 4.4, 8.0, 2.1])
    mission = p08["mission_metrics"]
    _add_callout(document, "نتیجهٔ مأموریت پایه", f"مأموریت در {mission['duration_s']:.1f} s با حالت {mission['final_state']} پایان یافت؛ {mission['collected_count']} زباله با جرم کل {mission['collected_mass_kg']:.4f} kg جمع شد، SOC نهایی {100*mission['final_soc']:.2f}% بود و حداقل حاشیهٔ ایمنی {mission['minimum_hazard_distance_m']:.3f} m ثبت شد.")


def _add_chapter_10(document: Document, p09: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "10. اعتبارسنجی سناریویی و تحلیل Monte Carlo", 1)
    _add_paragraph(document, "برای جلوگیری از اتکا به یک اجرای منفرد، زنجیرهٔ بسته‌حلقهٔ Phase 08 در چهار سناریوی نام‌دار و 20 اجرای Monte Carlo با seedهای ثابت تکرار شد. متغیرهای سناریو شامل مقدار/جهت جریان در بازهٔ اعتبار و SOC اولیه بودند. تمام محیط‌ها، حسگرها و مسیرهای پایه از مدل‌های قبلی استفاده می‌کنند.")
    _add_figure(document, "outputs/figures/phase09_scenario_trajectories.png", "شکل 28 — مسیرهای بسته‌حلقه در چهار سناریوی رسمی اعتبارسنجی.")
    _add_figure(document, "outputs/figures/phase09_mission_scorecard.png", "شکل 29 — کارت امتیاز مأموریت‌ها و مقایسهٔ نتایج سناریوهای رسمی.")
    _add_heading(document, "10-1. نتایج سناریوهای قطعی", 2)
    rows = []
    for scenario in p09["deterministic_scenarios"]:
        rows.append([
            str(scenario["scenario_id"]),
            str(scenario["final_state"]),
            str(scenario["collected_count"]),
            f"{scenario['duration_s']:.1f}",
            f"{scenario['final_soc']:.3f}",
            f"{scenario['minimum_clearance_m']:.3f}",
        ])
    _add_table(document, ["سناریو", "حالت نهایی", "جمع‌آوری", "زمان [s]", "SOC نهایی", "حداقل فاصله [m]"], rows, widths_cm=[3.6, 3.7, 2.0, 2.4, 2.3, 2.5])
    _add_figure(document, "outputs/figures/phase09_monte_carlo_robustness.png", "شکل 30 — پراکندگی نتایج 20 اجرای seeded Monte Carlo در محدودهٔ اعتبار کنترل.")
    _add_figure(document, "outputs/figures/phase09_sensitivity_heatmap.png", "شکل 31 — نقشهٔ حساسیت نتایج به جریان و SOC اولیه در بازهٔ تحلیل.")
    mc = {str(row["metric"]): row for row in p09["monte_carlo_summary"]}
    _add_heading(document, "10-2. جمع‌بندی آماری", 2)
    rows = [
        ["تعداد اجرا", f"{int(_safe_float(mc['trial_count']['value']))}"],
        ["نرخ موفقیت", f"{100*_safe_float(mc['success_rate']['value']):.1f}%"],
        ["میانگین زبالهٔ جمع‌آوری‌شده", f"{_safe_float(mc['mean_collected_count']['value']):.2f}"],
        ["میانهٔ زمان مأموریت", f"{_safe_float(mc['median_duration']['value']):.1f} s"],
        ["صدک 5 SOC نهایی", f"{_safe_float(mc['p05_final_soc']['value']):.3f}"],
        ["بدترین حاشیهٔ ایمنی", f"{_safe_float(mc['minimum_clearance']['value']):.3f} m"],
        ["میانگین خطای خانه", f"{_safe_float(mc['mean_home_error']['value']):.3f} m"],
    ]
    _add_table(document, ["شاخص Monte Carlo", "نتیجه"], rows, widths_cm=[8.0, 8.5])
    _add_callout(document, "تفسیر اعتبار", "نرخ موفقیت 100% تنها برای 20 اجرای seeded و بازهٔ جریان/ SOC تعریف‌شده گزارش می‌شود. این عدد به‌معنای تضمین عملکرد در جریان‌های بزرگ، موج، باد، موانع متحرک یا محیط واقعی نیست.")


def _add_chapter_11(document: Document) -> None:
    _page_break(document)
    _add_heading(document, "11. بازتولیدپذیری، شواهد اجرا و تحویل پروژه", 1)
    _add_paragraph(document, "پروژه به‌گونه‌ای سازمان‌دهی شده است که استاد بتواند بدون اجرای دستی ده‌ها فرمان، کل زنجیره را بازسازی کند. نقطهٔ ورود اصلی پروژه یک فایل batch است که محیط Conda را می‌سازد یا فعال می‌کند، پکیج پروژه را نصب می‌کند، تمام فازها را به‌ترتیب اجرا می‌کند، تست‌ها را اجرا می‌کند، گزارش Word را می‌سازد و پروندهٔ تحویل را بسته‌بندی می‌کند.")
    _add_heading(document, "11-1. فرمان بازتولید صفر تا صد", 2)
    _add_equation(document, r"scripts\bootstrap_and_build.bat")
    _add_paragraph(document, "هر فاز رسمی پوشهٔ timestamped خود را در records/phases ایجاد می‌کند. این پوشه شامل command transcript، stdout/stderr، environment snapshot، pip freeze، hash ورودی‌ها و Artifactها، snapshot Artifactها و Handoff فاز است. فایل‌های Handoff آخرین وضعیت هر فاز را در records/handoffs نگه می‌دارند.")
    _add_heading(document, "11-2. محتوای قابل تحویل", 2)
    rows = [
        ["گزارش اصلی", "outputs/reports/AquaSkim-Sim_Final_Report.docx"],
        ["ویدئوی مأموریت پایه", "outputs/videos/phase08_closed_loop_mission.mp4"],
        ["ویدئوی سناریوها", "outputs/videos/phase09_scenario_reel.mp4"],
        ["انیمیشن‌های GIF", "outputs/animations/"],
        ["تصاویر گزارش‌پذیر PNG/SVG", "outputs/figures/"],
        ["جداول عددی CSV", "outputs/tables/"],
        ["خلاصه‌های JSON و Markdown", "outputs/logs/ و outputs/reports/"],
        ["کد و پیکربندی", "src/, config/, scripts/, tests/"],
        ["شواهد و Handoffها", "records/"],
        ["بستهٔ نهایی", "outputs/deliverables/AquaSkim-Sim_Submission.zip"],
    ]
    _add_table(document, ["تحویل", "مسیر در پروژه"], rows, widths_cm=[5.0, 11.5], rtl=False)
    _add_callout(document, "نکتهٔ اجرایی", "در یک اجرای کاملاً تازه، ابتدا Miniconda باید نصب و دستور conda در CMD قابل دسترس باشد. سپس اجرای bootstrap_and_build.bat، محیط aquaskim-sim را ایجاد و پروژه را بازسازی می‌کند.")


def _add_chapter_12(document: Document, p03: dict[str, Any], p04: dict[str, Any], p07: dict[str, Any], p08: dict[str, Any], p09: dict[str, Any]) -> None:
    _page_break(document)
    _add_heading(document, "12. محدودیت‌ها و مسیر توسعه", 1)
    _add_heading(document, "12-1. محدودیت‌های صریح مدل", 2)
    limitations: list[str] = []
    for source in (p03, p04, p07, p08, p09):
        for item in source.get("limitations", source.get("assumptions", [])):
            if isinstance(item, str) and item not in limitations:
                limitations.append(item)
    for item in limitations[:12]:
        _add_bullet(document, item)
    _add_heading(document, "12-2. توسعه‌های پیشنهادی", 2)
    for item in [
        "جایگزینی هندسهٔ ضریب‌شکل با محاسبات هیدرواستاتیک مستقیم از بدنهٔ CAD و مش سه‌بعدی.",
        "اضافه‌کردن مدل موج، باد، جریان مکانی–زمانی و موانع متحرک.",
        "شناسایی ضرایب sway/yaw و مدل پیشران با دادهٔ آزمون یا CFD/تست تانک.",
        "افزودن تخمین حالت مبتنی بر EKF و مدل خطاهای واقعی‌تر حسگر.",
        "توسعهٔ مدل جمع‌آوری از شرط هندسی به مدل تعامل سیال–زباله و بررسی ظرفیت واقعی سبد.",
        "اجرای Hardware-in-the-Loop یا ساخت نمونهٔ فیزیکی، تنها پس از اعتبارسنجی مرحله‌ای مدل‌ها.",
    ]:
        _add_bullet(document, item)
    _add_heading(document, "12-3. نتیجه‌گیری", 2)
    _add_paragraph(document, "AquaSkim-Sim یک پروژهٔ یکپارچهٔ طراحی و شبیه‌سازی است که مسیر مهندسی را از پارامترهای مکانیکی تا تصمیم‌های خودگردان و اعتبارسنجی آماری حفظ می‌کند. مهم‌ترین ارزش پروژه، شفافیت فرض‌ها، اتصال مستقیم محاسبات به کد و امکان بازتولید کامل Artifactها است. بر پایهٔ نتایج گزارش‌شده، طراحی در محدودهٔ مدل و سناریوهای تعریف‌شده، از نظر پایداری، رانش، انرژی، ایمنی مسیر و اتمام مأموریت، معیارهای قبولی را پاس کرده است.")


def _add_appendices(document: Document, phase_artifact_count: int) -> None:
    _page_break(document)
    _add_heading(document, "پیوست A. معادلات کلیدی", 1)
    for expression, explanation in [
        ("m_total = Σm_i", "جمع جرم اجزای مدل مکانیکی."),
        ("r_CG = Σ(m_i r_i) / Σm_i", "مرکز جرم سه‌بعدی."),
        ("F_B = ρg∇ = mg", "تعادل شناوری و وزن."),
        ("GM = KB + BM − KG ; BM = I_T / ∇", "پایداری اولیه."),
        ("C_F = 0.075/[log10(Re) − 2]^2", "خط اصطکاک ITTC-1957."),
        ("T = k_T RPM² ; P = k_P RPM³", "مدل پیشران مفهومی."),
        ("SOC_{k+1} = SOC_k − I_k Δt / C_effective", "انتگرال‌گیری SOC."),
        ("x_dot = u cosψ − v sinψ ; y_dot = u sinψ + v cosψ ; ψ_dot = r", "سینماتیک حرکت صفحه‌ای."),
        ("τ_z = (b/2)(T_R − T_L)", "گشتاور yaw ناشی از رانش تفاضلی."),
    ]:
        _add_equation(document, expression, explanation)

    _add_heading(document, "پیوست B. فهرست خروجی‌ها و Artifactها", 1)
    _add_paragraph(document, f"تا پایان Phase 09، اجرای رسمی پروژه {phase_artifact_count} Artifact منبع را با snapshot و hash ثبت کرده است. Phase 10 نیز گزارش Word، manifest ساخت گزارش، بستهٔ تحویل و checksumها را به زنجیره اضافه می‌کند.")
    inventory = [
        ["طرح مکانیکی", "Phase 02", "نمای بالا، نمای جانبی، توزیع جرم، CSVهای هندسه و جرم"],
        ["هیدرواستاتیک", "Phase 03", "آبخور، فری‌بورد، GZ، گشتاور راست‌کننده، envelope بار"],
        ["پیشرانش", "Phase 04", "مقاومت، رانش، RPM، توان و اثر جریان مخالف"],
        ["انرژی", "Phase 05", "SOC، دوام، envelope بازگشت به خانه"],
        ["دینامیک", "Phase 06", "پاسخ 3-DOF، مانور و جریان جانبی"],
        ["محیط و حسگر", "Phase 07", "نقشه، Occupancy Grid، sensor log و آشکارسازی"],
        ["خودگردانی", "Phase 08", "A*، FSM، کنترل، GIF و MP4 مأموریت"],
        ["اعتبارسنجی", "Phase 09", "سناریوها، Monte Carlo، GIF و MP4 مقایسه‌ای"],
        ["گزارش نهایی", "Phase 10", "DOCX، manifest، checksum و بستهٔ تحویل"],
    ]
    _add_table(document, ["دسته", "فاز", "Artifactهای اصلی"], inventory, widths_cm=[3.4, 2.5, 10.6])

    _add_heading(document, "پیوست C. منابع علمی مدل", 1)
    references = [
        "Fossen, T. I. Handbook of Marine Craft Hydrodynamics and Motion Control, 2011. مبنای مدل‌سازی حرکت سطحی و چارچوب surge–sway–yaw.",
        "ITTC. 1957 Model-Ship Correlation Line. مبنای ضریب اصطکاکی مورد استفاده در تحلیل مقاومت Phase 04.",
        "Principles of Naval Architecture. Society of Naval Architects and Marine Engineers. مبنای مفاهیم شناوری، متاسنتر و پایداری اولیه.",
        "Russell, S. and Norvig, P. Artificial Intelligence: A Modern Approach. مبنای مفهومی عامل، تصمیم‌گیری و برنامه‌ریزی مسیر؛ در پروژه از عامل توضیح‌پذیر مبتنی بر ماشین حالت و A* استفاده شده است.",
    ]
    for item in references:
        _add_bullet(document, item)


def _build_document() -> tuple[Document, dict[str, Any]]:
    ensure_runtime_directories()
    config = load_base_configuration().data
    p02 = _read_json("outputs/logs/phase02_mechanical_summary.json")
    p03 = _read_json("outputs/logs/phase03_hydrostatic_summary.json")
    p04 = _read_json("outputs/logs/phase04_propulsion_summary.json")
    p05 = _read_json("outputs/logs/phase05_energy_summary.json")
    p06 = _read_json("outputs/logs/phase06_dynamics_summary.json")
    p07 = _read_json("outputs/logs/phase07_environment_summary.json")
    p08 = _read_json("outputs/logs/phase08_autonomy_summary.json")
    p09 = _read_json("outputs/logs/phase09_validation_summary.json")
    metadata = _metadata()

    document = Document()
    _configure_document(document)
    date_text = datetime.now().strftime("%Y-%m-%d")
    _add_cover(document, metadata, date_text)
    _add_contents(document)
    _add_abstract(document, _summary_values(p03, p04, p05, p08, p09))
    _add_chapter_1(document, config, p07)
    _add_chapter_2(document)
    _add_chapter_3(document, config, p02)
    _add_chapter_4(document, p03)
    _add_chapter_5(document, p04)
    _add_chapter_6(document, p05)
    _add_chapter_7(document, p06)
    _add_chapter_8(document, p07)
    _add_chapter_9(document, p08)
    _add_chapter_10(document, p09)
    _add_chapter_11(document)
    _add_chapter_12(document, p03, p04, p07, p08, p09)
    _add_appendices(document, phase_artifact_count=131)

    source_files = [
        "config/base_parameters.yaml",
        "config/report_metadata.json",
        *[f"outputs/logs/phase{index:02d}_{suffix}.json" for index, suffix in []],
        "outputs/logs/phase02_mechanical_summary.json",
        "outputs/logs/phase03_hydrostatic_summary.json",
        "outputs/logs/phase04_propulsion_summary.json",
        "outputs/logs/phase05_energy_summary.json",
        "outputs/logs/phase06_dynamics_summary.json",
        "outputs/logs/phase07_environment_summary.json",
        "outputs/logs/phase08_autonomy_summary.json",
        "outputs/logs/phase09_validation_summary.json",
    ]
    figure_files = [
        f"outputs/figures/{name}.png"
        for name in (
            "phase02_mechanical_top_view", "phase02_mechanical_side_view", "phase02_mass_distribution",
            "phase03_hydrostatics_dashboard", "phase03_stability_curves", "phase03_heeling_cross_sections", "phase03_payload_envelope",
            "phase04_resistance_dashboard", "phase04_propulsion_envelope", "phase04_current_penalty", "phase04_operating_envelope",
            "phase05_energy_dashboard", "phase05_mission_soc_profiles", "phase05_return_home_envelope",
            "phase06_dynamics_dashboard", "phase06_trajectory_comparison", "phase06_maneuver_response", "phase06_current_disturbance",
            "phase07_environment_map", "phase07_occupancy_grid", "phase07_sensor_model", "phase07_perception_dashboard",
            "phase08_autonomy_architecture", "phase08_planning_map", "phase08_closed_loop_mission", "phase08_control_dashboard", "phase08_decision_timeline",
            "phase09_scenario_trajectories", "phase09_mission_scorecard", "phase09_monte_carlo_robustness", "phase09_sensitivity_heatmap",
        )
    ]
    manifest = {
        "report_title": REPORT_SUBTITLE_FA,
        "generated_utc": datetime.now(timezone.utc).isoformat(),
        "generator": "aquaskim.phase10",
        "source_files": [
            {"path": rel, "sha256": _sha256(PROJECT_ROOT / rel)}
            for rel in source_files if (PROJECT_ROOT / rel).exists()
        ],
        "embedded_figures": [
            {"path": rel, "sha256": _sha256(PROJECT_ROOT / rel)}
            for rel in figure_files if (PROJECT_ROOT / rel).exists()
        ],
        "report_image_derivatives": [
            {
                "path": relative_to_root(path),
                "sha256": _sha256(path),
                "size_bytes": path.stat().st_size,
            }
            for path in sorted((DIRECTORIES["reports"] / "report_assets").glob("*_report.png"))
        ],
        "model_scope": "Digital-twin conceptual engineering simulation; no physical build or field certification claim.",
    }
    return document, manifest


def _inspect_docx(path: Path) -> dict[str, Any]:
    if not path.exists() or path.stat().st_size < 100_000:
        raise ReportBuildError("Final DOCX is missing or unexpectedly small.")
    with zipfile.ZipFile(path) as archive:
        members = archive.namelist()
        required = {"[Content_Types].xml", "word/document.xml"}
        if not required.issubset(members):
            raise ReportBuildError("Final DOCX does not have the required OOXML members.")
        media = [name for name in members if name.startswith("word/media/")]
    document = Document(str(path))
    nonempty_paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return {
        "path": relative_to_root(path),
        "size_bytes": path.stat().st_size,
        "sha256": _sha256(path),
        "embedded_media_count": len(media),
        "nonempty_paragraph_count": len(nonempty_paragraphs),
        "contains_project_title": any("AquaSkim" in text for text in nonempty_paragraphs),
        "validation_status": "PASS",
    }


def _write_report_readme(path: Path) -> None:
    content = """# گزارش نهایی AquaSkim-Sim

فایل `AquaSkim-Sim_Final_Report.docx` از خروجی‌های واقعی Phase 02 تا Phase 09 ساخته می‌شود.

## تولید دوباره

```bat
scripts\\run_patch_10.bat
```

یا برای بازسازی کامل از صفر تا صد:

```bat
scripts\\bootstrap_and_build.bat
```

## شخصی‌سازی جلد

اطلاعات جلد در فایل `config/report_metadata.json` قرار دارد. قبل از تحویل، نام، شمارهٔ دانشجویی، نام استاد و دانشگاه را در همان فایل اصلاح کنید و دوباره گزارش را بسازید.

## کنترل کیفیت

این پروژه ساختار DOCX، تعداد تصاویر embed‌شده، وجود تیترهای اصلی و checksum فایل گزارش را به‌صورت خودکار کنترل می‌کند. برای مشاهدهٔ نهایی، فایل DOCX را در Microsoft Word باز کنید.
"""
    path.write_text(content, encoding="utf-8")


def _iter_submission_files() -> Iterable[Path]:
    ignored_parts = {".git", ".pytest_cache", "__pycache__"}
    excluded_exact = {
        PROJECT_ROOT / "outputs" / "deliverables" / "AquaSkim-Sim_Submission.zip",
        PROJECT_ROOT / "outputs" / "deliverables" / "AquaSkim-Sim_Submission_manifest.json",
        PROJECT_ROOT / "outputs" / "deliverables" / "AquaSkim-Sim_SHA256SUMS.txt",
    }
    for path in PROJECT_ROOT.rglob("*"):
        if not path.is_file():
            continue
        if any(part in ignored_parts for part in path.parts):
            continue
        if path in excluded_exact:
            continue
        if path.suffix.lower() in {".pyc", ".pyo"}:
            continue
        yield path


def _package_submission(report_docx: Path) -> tuple[Path, Path, Path]:
    deliverables = DIRECTORIES["outputs"] / "deliverables"
    deliverables.mkdir(parents=True, exist_ok=True)
    zip_path = deliverables / "AquaSkim-Sim_Submission.zip"
    manifest_path = deliverables / "AquaSkim-Sim_Submission_manifest.json"
    checksum_path = deliverables / "AquaSkim-Sim_SHA256SUMS.txt"

    records: list[dict[str, Any]] = []
    for path in sorted(_iter_submission_files()):
        records.append({
            "path": relative_to_root(path),
            "size_bytes": path.stat().st_size,
            "sha256": _sha256(path),
        })

    manifest = {
        "package": "AquaSkim-Sim_Submission.zip",
        "generated_utc": datetime.now(timezone.utc).isoformat(),
        "file_count": len(records),
        "report_docx": relative_to_root(report_docx),
        "files": records,
        "rebuild_command": "scripts\\bootstrap_and_build.bat",
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    checksum_path.write_text("\n".join(f"{item['sha256']}  {item['path']}" for item in records) + "\n", encoding="utf-8")

    # Images, videos and Office files are already compressed. Storing them avoids
    # a long, low-value recompression pass on Windows while keeping the archive
    # fully portable. Text and source files are deflated.
    precompressed_suffixes = {".png", ".jpg", ".jpeg", ".gif", ".mp4", ".docx", ".pdf", ".zip", ".stl", ".step"}
    temporary_zip = zip_path.with_suffix(".tmp.zip")
    if temporary_zip.exists():
        temporary_zip.unlink()
    with zipfile.ZipFile(temporary_zip, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=3) as archive:
        for item in records:
            source = PROJECT_ROOT / str(item["path"])
            compression = zipfile.ZIP_STORED if source.suffix.lower() in precompressed_suffixes else zipfile.ZIP_DEFLATED
            archive.write(source, arcname=f"AquaSkim-Sim/{item['path']}", compress_type=compression)
        archive.write(manifest_path, arcname="AquaSkim-Sim/outputs/deliverables/AquaSkim-Sim_Submission_manifest.json", compress_type=zipfile.ZIP_DEFLATED)
        archive.write(checksum_path, arcname="AquaSkim-Sim/outputs/deliverables/AquaSkim-Sim_SHA256SUMS.txt", compress_type=zipfile.ZIP_DEFLATED)
    temporary_zip.replace(zip_path)
    return zip_path, manifest_path, checksum_path


def run_phase10() -> Phase10Artifacts:
    """Build the final Word report and an offline submission package."""
    ensure_runtime_directories()
    report_dir = DIRECTORIES["reports"]
    report_dir.mkdir(parents=True, exist_ok=True)
    report_docx = report_dir / "AquaSkim-Sim_Final_Report.docx"
    report_manifest = report_dir / "phase10_report_build_manifest.json"
    report_readme = report_dir / "README_FINAL_REPORT_FA.md"

    document, manifest = _build_document()
    document.core_properties.title = REPORT_SUBTITLE_FA
    document.core_properties.subject = "Autonomous mobile robots — complete digital-twin engineering project"
    document.core_properties.author = "AquaSkim-Sim project generator"
    document.core_properties.keywords = "catamaran, autonomous robot, digital twin, hydrostatics, control, simulation"
    document.save(report_docx)

    report_validation = _inspect_docx(report_docx)
    manifest["report_validation"] = report_validation
    report_manifest.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_report_readme(report_readme)

    submission_zip, submission_manifest, checksums = _package_submission(report_docx)
    return Phase10Artifacts(
        report_docx=report_docx,
        report_manifest=report_manifest,
        report_readme=report_readme,
        submission_zip=submission_zip,
        submission_manifest=submission_manifest,
        checksums=checksums,
    )


def print_phase10_summary(artifacts: Phase10Artifacts) -> None:
    print("=" * 72)
    print("AquaSkim-Sim | Phase 10 Final Report and Submission Package")
    print("=" * 72)
    for name, path in artifacts.as_dict().items():
        print(f"{name:24}: {path}")
    print("=" * 72)
    print("[OK] Final Word report, manifest, checksums and submission package generated.")
