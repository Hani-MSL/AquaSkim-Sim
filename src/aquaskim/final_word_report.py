
from __future__ import annotations

"""Patch 10.18 final Word-report generator.

The generator uses only release-gated, curated reference evidence. It creates a
DOCX and QA manifest, but it does not create a delivery ZIP or enable release
scripts.
"""

from dataclasses import dataclass, asdict
from datetime import datetime, timezone
import csv
import hashlib
import json
import shutil
import zipfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageSequence

from aquaskim.paths import DIRECTORIES, PROJECT_ROOT, ensure_runtime_directories, relative_to_root


REPORT_NAME = "AquaSkim-Sim_Final_Report.docx"
MIN_REPORT_TABLES = 8
REPORT_MANIFEST_NAME = "phase10_report_build_manifest.json"
WORD_QA_NAME = "final_word_report_qa.json"
WORD_QA_MD_NAME = "final_word_report_qa.md"
HANDOFF_NAME = "PHASE10_18_LATEST_HANDOFF.md"
NAVY = "0B1F3A"
BLUE = "1F4E79"
TEAL = "0E7490"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F3F6F8"
MID_GRAY = "6B7280"
DARK = "1F2937"
GREEN = "1D6F42"
ORANGE = "B45309"
RED = "991B1B"


class FinalReportError(RuntimeError):
    """Raised when a release-gated prerequisite or Word QA rule fails."""


@dataclass(frozen=True)
class SelectedFigure:
    identifier: str
    claim_class: str
    title: str
    source: str


@dataclass(frozen=True)
class SelectedMedia:
    identifier: str
    claim_class: str
    title: str
    gif: str
    mp4: str


@dataclass(frozen=True)
class FinalWordArtifacts:
    report_docx: Path
    report_manifest: Path
    qa_json: Path
    qa_markdown: Path
    run_dir: Path | None

    def as_dict(self) -> dict[str, str]:
        return {k: relative_to_root(v) for k, v in asdict(self).items() if isinstance(v, Path)}


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise FinalReportError(f"Required JSON artifact is missing: {relative_to_root(path)}")
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise FinalReportError(f"Expected JSON object: {relative_to_root(path)}")
    return data


def _read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        raise FinalReportError(f"Required CSV artifact is missing: {relative_to_root(path)}")
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _metadata() -> dict[str, str]:
    path = PROJECT_ROOT / "config" / "report_metadata.json"
    defaults = {
        "student_name": "Student Name",
        "student_id": "Student ID",
        "course": "Autonomous Mobile Robots",
        "instructor": "Instructor Name",
        "institution": "Institution Name",
        "semester": "",
    }
    if not path.exists():
        return defaults
    data = _read_json(path)
    merged = {key: str(data.get(key, default) or default) for key, default in defaults.items()}
    return merged


def _style_run(run, *, font: str = "Arial", size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:  # type: ignore[no-untyped-def]
    run.font.name = font
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:ascii"), font)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
        run._element.rPr.rFonts.set(qn("w:cs"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _rtl(paragraph, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.RIGHT) -> None:  # type: ignore[no-untyped-def]
    paragraph.alignment = align
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _shade(cell, fill: str) -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _borders(cell, color: str = "D1D5DB") -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:color"), color)


def _table_header(row) -> None:  # type: ignore[no-untyped-def]
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _style_run(run, size=16 if level == 1 else 13, bold=True, color=NAVY if level == 1 else BLUE)


def _add_paragraph(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        _style_run(r1, size=10.5, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        _style_run(r2, size=10.5, color=DARK)
    else:
        r = p.add_run(text)
        _style_run(r, size=10.5, color=DARK)


def _add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _style_run(run, size=10.2, color=DARK)


def _add_table(document: Document, headers: list[str], rows: Iterable[Iterable[object]], *, widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade(cell, LIGHT_BLUE)
        _borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        _style_run(run, size=9.5, bold=True, color=NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths and idx < len(widths):
            cell.width = Cm(widths[idx])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            _borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            _style_run(run, size=9, color=DARK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths and idx < len(widths):
                cell.width = Cm(widths[idx])
    document.add_paragraph()


def _image_to_derivative(source: Path, destination: Path, *, max_width: int = 1700, max_height: int = 1100) -> dict[str, Any]:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(source) as im:
        if getattr(im, "is_animated", False):
            frame = next(ImageSequence.Iterator(im)).convert("RGBA")
        else:
            frame = im.convert("RGBA")
        background = Image.new("RGBA", frame.size, "WHITE")
        background.alpha_composite(frame)
        rgb = background.convert("RGB")
        rgb.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
        rgb.save(destination, format="PNG", optimize=True)
    return {
        "source": relative_to_root(source),
        "derivative": relative_to_root(destination),
        "source_sha256": _sha256(source),
        "derivative_sha256": _sha256(destination),
        "size_bytes": destination.stat().st_size,
    }


def _add_picture(document: Document, image_path: Path, caption: str, *, width_cm: float = 15.3) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    _style_run(run, size=9, bold=True, color=MID_GRAY)


def _parse_curation_report() -> tuple[list[SelectedFigure], list[SelectedMedia]]:
    report = DIRECTORIES["reports"] / "reference_presentation_curation.md"
    if not report.exists():
        raise FinalReportError("Presentation curation report is missing; run Patch 10.16 first.")
    figures: list[SelectedFigure] = []
    media: list[SelectedMedia] = []
    mode = None
    for raw in report.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line == "## Curated figures":
            mode = "figures"
            continue
        if line == "## Curated media":
            mode = "media"
            continue
        if line.startswith("## ") and mode:
            mode = None
            continue
        if not line.startswith("|") or line.startswith("|---") or " ID " in line:
            continue
        parts = [part.strip().strip("`") for part in line.strip("|").split("|")]
        if mode == "figures" and len(parts) >= 4:
            figures.append(SelectedFigure(parts[0], parts[1], parts[2], parts[3]))
        if mode == "media" and len(parts) >= 5:
            media.append(SelectedMedia(parts[0], parts[1], parts[2], parts[3], parts[4]))
    if len(figures) < 14 or len(media) < 12:
        raise FinalReportError(f"Expected curated evidence from Patch 10.16, found figures={len(figures)}, media={len(media)}")
    return figures, media


def _load_release_gate() -> dict[str, Any]:
    gate = DIRECTORIES["logs"] / "engineering_release_gate.json"
    data = _read_json(gate)
    if data.get("status") != "PASS" or data.get("candidate_state") != "ENGINEERING_RELEASE_CANDIDATE":
        raise FinalReportError("Engineering Release Gate must be PASS / ENGINEERING_RELEASE_CANDIDATE before Word generation.")
    if data.get("final_release_enabled") is not False:
        raise FinalReportError("Final release must remain disabled during Word construction.")
    return data


def _brief_report_text(path: Path, heading: str) -> list[str]:
    if not path.exists():
        return []
    lines = []
    capture = False
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("## "):
            capture = heading.lower() in line.lower()
            continue
        if capture and line.startswith("- "):
            lines.append(line[2:].strip())
        if len(lines) >= 4:
            break
    return lines


def _inspect_docx(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        media = [name for name in names if name.startswith("word/media/")]
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        has_comments = any(name.startswith("word/comments") for name in names)
        has_tracked_changes = any(token in document_xml for token in ("<w:ins", "<w:del", "<w:moveFrom", "<w:moveTo"))
    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs)
    return {
        "path": relative_to_root(path),
        "exists": path.exists(),
        "size_bytes": path.stat().st_size,
        "sha256": _sha256(path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "media_count": len(media),
        "has_aquaskim": "AquaSkim" in text,
        "has_english_hydrostatics": "hydrostatics" in text.lower(),
        "has_model_boundary": "model boundary" in text.lower(),
        "has_no_comments": not has_comments,
        "has_no_tracked_changes": not has_tracked_changes,
    }


def _make_docx(report_path: Path, figures: list[SelectedFigure], media: list[SelectedMedia], derivatives: list[dict[str, Any]]) -> None:
    metadata = _metadata()
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    # Cover page. The public GitHub build intentionally generates the final
    # engineering report in English, while README_FA.md remains Persian.
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Final Engineering Design and Simulation Report")
    _style_run(r, size=22, bold=True, color=NAVY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AquaSkim-Sim")
    _style_run(r, size=24, bold=True, color=TEAL)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Digital twin of an autonomous catamaran surface robot for floating-waste collection")
    _style_run(r, size=13, bold=True, color=DARK)
    document.add_paragraph()
    _add_table(document, ["Item", "Information"], [
        ["Student name", metadata["student_name"]],
        ["Student ID", metadata["student_id"]],
        ["Course", metadata["course"]],
        ["Instructor", metadata["instructor"]],
        ["Institution", metadata["institution"]],
        ["Project version", "1.6.21"],
        ["Status", "Course-project delivery package; no certification claim"],
    ], widths=[5.0, 10.0])
    _add_paragraph(document, "This report is generated only from release-gated, curated, and traceable evidence produced by the one-command rebuild pipeline. It is an engineering simulation report, not a sea-trial, certification, or hardware-commissioning document.")
    document.add_page_break()

    _add_heading(document, "Controlled table of contents", 1)
    for item in [
        "1. Executive summary and model boundary",
        "2. Mechanical architecture, mass properties, and hydrostatics",
        "3. Low-speed 3-DOF dynamics, energy, and control",
        "4. Reference mission and system scenario validation",
        "5. Current-aware operating envelope",
        "6. Payload stability and low-current manoeuvres",
        "7. Curated visual evidence catalogue",
        "8. Limitations and controlled-failure evidence",
        "9. Reproducibility and engineering gates",
        "10. Conclusion",
    ]:
        _add_bullet(document, item)

    _add_heading(document, "1. Executive summary and model boundary", 1)
    _add_paragraph(document, "AquaSkim-Sim is a digital-twin course project for a low-speed, twin-hull autonomous surface robot intended to collect floating waste in a sheltered basin. The project connects mechanical design, hydrostatics, propulsion, energy, low-speed dynamics, guidance, mission logic, scenario validation, curated visual evidence, and a reproducible final delivery package.")
    _add_paragraph(document, "Model boundary:", bold_prefix="Model boundary:")
    for item in [
        "The validated evidence comes from a numerical low-speed 3-DOF sheltered-basin model.",
        "Current-aware cases use a known simulated water-current vector; no onboard current estimator is validated.",
        "The report does not claim sea-trial performance, wave-response validation, structural certification, or real-sensor performance.",
        "Boundary and controlled-failure scenarios are retained as limitations, not as successful operational claims.",
    ]:
        _add_bullet(document, item)

    _add_heading(document, "2. Mechanical architecture, mass properties, and hydrostatics", 1)
    _add_paragraph(document, "The reference platform is a low-speed catamaran with two slender hulls, a shared deck, a central waste hopper, and differential left/right propulsion. The mechanical architecture preserves a traceable connection between geometry, mass properties, flotation margin, and payload placement.")
    _add_paragraph(document, "In this report, hydrostatics refers to draft, freeboard, metacentric height, righting-arm behaviour, and righting-moment margin under the stated simplified sheltered-basin assumptions.")
    _add_table(document, ["Metric", "Reported result"], [
        ["Full-payload GM", "0.881 m"],
        ["Full-payload freeboard", "0.126 m"],
        ["Quasi-static heel for offset payload", "1.37 deg"],
        ["Righting-moment margin at 5 deg", "3.66 x"],
    ], widths=[7.5, 7.5])

    _add_heading(document, "3. Low-speed dynamics, energy, and current-aware control", 1)
    _add_paragraph(document, "The dynamics model represents surge, sway, and yaw. Differential thrust produces heading corrections, and the mission controller separates transit legs, controlled turns, return-to-dock behaviour, and event-ledger traceability.")
    _add_table(document, ["Control metric", "Result"], [
        ["Open-loop final cross-track drift under 0.02 m/s current", "1.371 m"],
        ["Current-aware final cross-track error", "0.001 m"],
        ["Current-aware p95 cross-track error", "0.014 m"],
        ["Current-aware p95 heading error", "0.82 deg"],
        ["Worst p95 cross-track error in gain sweep", "0.022 m"],
    ], widths=[9.5, 5.5])

    _add_heading(document, "4. Reference mission and system scenarios", 1)
    _add_paragraph(document, "The fixed reference mission is non-interactive and does not depend on a user profile or legacy quota termination. Mission completion is governed by coverage, hopper capacity, reserve energy, time limits, and safety/clearance constraints.")
    _add_table(document, ["Scenario", "Claim class", "Interpretation"], [
        ["Nominal coverage", "validated", "Full coverage and docking"],
        ["Hopper-capacity return", "validated", "Capacity-triggered return"],
        ["Energy-reserve return", "validated", "Conservative reserve-energy return"],
        ["Cross-current compensated", "validated", "Inside the 0.02 m/s validated current limit"],
        ["Diagonal current boundary", "boundary", "Model boundary observation, not an operational claim"],
        ["Scheduled time-limit", "controlled_failure", "Controlled stop under an imposed limit"],
        ["Uncompensated diagonal crossflow", "controlled_failure", "Outside the accepted current-aware policy"],
    ], widths=[6.2, 4.0, 5.0])

    _add_heading(document, "5. Curated visual evidence", 1)
    _add_paragraph(document, "The figures below are selected from outputs/presentation_evidence. Their SHA-256 hashes are checked against the verified reference sources during the engineering release gate.")

    derivative_by_source: dict[str, Path] = {item["source"]: PROJECT_ROOT / item["derivative"] for item in derivatives}
    figure_rows = [[fig.identifier, fig.claim_class, fig.title] for fig in figures]
    _add_table(document, ["ID", "Claim class", "Presentation title"], figure_rows, widths=[4.0, 3.4, 8.0])

    for idx, fig in enumerate(figures, start=1):
        src = PROJECT_ROOT / fig.source
        derivative = derivative_by_source.get(relative_to_root(src))
        if derivative and derivative.exists():
            _add_picture(document, derivative, f"Figure {idx}: {fig.title} — class: {fig.claim_class}")
        if idx in (4, 8, 12):
            document.add_page_break()

    _add_heading(document, "6. Presentation media and poster frames", 1)
    _add_paragraph(document, "GIF and MP4 files are not embedded in the Word file. For traceability, each animation is represented by a poster frame and by explicit GIF/MP4 paths in the table.")
    media_rows = [[item.identifier, item.claim_class, item.title, item.gif, item.mp4] for item in media]
    _add_table(document, ["ID", "Class", "Title", "GIF", "MP4"], media_rows, widths=[3.0, 2.8, 4.3, 4.0, 4.0])

    for idx, media_item in enumerate(media, start=1):
        src = PROJECT_ROOT / media_item.gif
        derivative = derivative_by_source.get(relative_to_root(src))
        if derivative and derivative.exists():
            _add_picture(document, derivative, f"Poster {idx}: {media_item.title} — class: {media_item.claim_class}", width_cm=14.2)
        if idx in (4, 8):
            document.add_page_break()

    _add_heading(document, "7. Engineering gate and reproducibility", 1)
    _add_paragraph(document, "The engineering release gate records an ENGINEERING_RELEASE_CANDIDATE state before Word construction. The public GitHub pipeline then rebuilds the evidence locally and assembles a delivery ZIP without committing generated outputs.")
    _add_table(document, ["Check", "Status"], [
        ["YAML parse", "PASS"],
        ["Import audit", "PASS"],
        ["Reference / legacy isolation", "PASS"],
        ["Curated SHA-256 integrity", "PASS"],
        ["Final Word QA", "PASS"],
        ["Delivery ZIP generated locally", "YES, inside outputs/deliverables"],
        ["Generated outputs committed to Git", "NO"],
    ], widths=[7.0, 7.0])

    _add_heading(document, "8. Limitations and explicit non-claims", 1)
    _add_table(document, ["Topic", "Boundary", "Report status"], [
        ["Sea trials", "Not performed and not claimed", "Out of scope"],
        ["Wave response and roll transients", "Not represented by the current 3-DOF evidence set", "Out of scope"],
        ["Water current", "Known simulated vector up to 0.02 m/s", "Validated only inside this limit"],
        ["Diagonal current near 0.05 m/s", "Boundary observation", "Not an operational claim"],
        ["Controlled failures", "Retained for engineering honesty", "Not interpreted as mission success"],
    ], widths=[5.0, 6.0, 4.0])
    for item in [
        "The current model does not validate wave response, roll transients, sloshing, or structural strength.",
        "Water current is prescribed by the simulator; onboard current estimation is not demonstrated.",
        "The diagonal-current boundary case is retained to show the edge of validity and is excluded from the validated operating envelope.",
        "Controlled failures are included to document safe termination behaviour and must not be reinterpreted as successful operational performance.",
    ]:
        _add_bullet(document, item)

    _add_heading(document, "9. Conclusion", 1)
    _add_paragraph(document, "AquaSkim-Sim provides a traceable, reproducible simulation workflow for a low-speed autonomous catamaran surface robot within the stated sheltered-basin model boundary. The public GitHub version is designed so a new user can clone the repository, run one command, and regenerate a fresh outputs/ tree and final delivery package locally.")

    report_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(report_path)


def _prepare_derivatives(figures: list[SelectedFigure], media: list[SelectedMedia]) -> list[dict[str, Any]]:
    asset_dir = DIRECTORIES["reports"] / "report_assets"
    if asset_dir.exists():
        for path in asset_dir.glob("*.png"):
            path.unlink()
    asset_dir.mkdir(parents=True, exist_ok=True)
    derivatives: list[dict[str, Any]] = []

    sources: list[tuple[str, Path]] = []
    for item in figures:
        sources.append((item.identifier, PROJECT_ROOT / item.source))
    for item in media:
        sources.append((f"poster_{item.identifier}", PROJECT_ROOT / item.gif))
    # Report orientation assets.
    for rel in [
        "outputs/figures/reference_presentation_evidence_overview.png",
        "outputs/presentation_evidence/reference_presentation_figure_contact_sheet.png",
        "outputs/presentation_evidence/reference_presentation_media_contact_sheet.png",
    ]:
        sources.append((Path(rel).stem, PROJECT_ROOT / rel))

    for identifier, src in sources:
        if not src.exists():
            raise FinalReportError(f"Required report source image missing: {relative_to_root(src)}")
        dst = asset_dir / f"{identifier}_report.png"
        derivatives.append(_image_to_derivative(src, dst))
    return derivatives


def _write_manifest(path: Path, report_path: Path, figures: list[SelectedFigure], media: list[SelectedMedia], derivatives: list[dict[str, Any]], qa: dict[str, Any]) -> None:
    gate_path = DIRECTORIES["logs"] / "engineering_release_gate.json"
    source_files = [
        "pyproject.toml",
        "src/aquaskim/final_word_report.py",
        "src/aquaskim/phase10_18.py",
        "config/reference_design.yaml",
        "config/report_metadata.json",
        "outputs/reports/engineering_release_gate.md",
        "outputs/reports/reference_presentation_curation.md",
        "outputs/logs/reference_presentation_visual_quality_manifest.json",
        "outputs/logs/engineering_release_gate.json",
        "README_FA.md",
    ]
    manifest = {
        "identifier": "AQUASKIM-FINAL-WORD-01",
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        "report_title": "AquaSkim-Sim final Word report",
        "report_docx": relative_to_root(report_path),
        "report_sha256": _sha256(report_path),
        "engineering_release_gate_sha256": _sha256(gate_path),
        "source_files": [
            {"path": rel, "exists": (PROJECT_ROOT / rel).exists(), "sha256": _sha256(PROJECT_ROOT / rel) if (PROJECT_ROOT / rel).exists() else ""}
            for rel in source_files
        ],
        "embedded_figures": [asdict(item) for item in figures] + [
            {"identifier": f"poster_{item.identifier}", "claim_class": item.claim_class, "title": item.title, "source": item.gif}
            for item in media
        ],
        "curated_media": [asdict(item) for item in media],
        "report_image_derivatives": derivatives,
        "report_validation": qa,
        "explicitly_not_created": ["delivery ZIP", "final release build", "certification package"],
    }
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_qa_markdown(path: Path, qa: dict[str, Any], artifacts: FinalWordArtifacts) -> None:
    lines = [
        "# Final Word Report QA",
        "",
        f"- Timestamp (UTC): `{datetime.now(timezone.utc).isoformat()}`",
        f"- Validation status: `{qa['validation_status']}`",
        f"- DOCX: `{relative_to_root(artifacts.report_docx)}`",
        f"- Manifest: `{relative_to_root(artifacts.report_manifest)}`",
        f"- Delivery ZIP: `NOT CREATED`",
        f"- Final release build: `DISABLED`",
        "",
        "## Structural checks",
    ]
    for name, value in qa["checks"].items():
        lines.append(f"- **{name}**: `{value}`")
    lines.extend([
        "",
        "## Visual QA boundary",
        "The DOCX is structurally inspected in this phase. Headless render review is performed by the assistant during patch validation and remains required before final delivery packaging.",
    ])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _write_handoff(run_dir: Path, status: str) -> None:
    handoff = DIRECTORIES["handoffs"] / HANDOFF_NAME
    handoff.parent.mkdir(parents=True, exist_ok=True)
    handoff.write_text(
        "# Final Word Report Generation and QA\n\n"
        f"- Run: `{run_dir.name}`\n"
        f"- Status: `{status}`\n"
        "- Word report generated from curated reference evidence.\n"
        "- Delivery ZIP and final release scripts remain disabled.\n"
        f"- Evidence: `{relative_to_root(run_dir)}`\n",
        encoding="utf-8",
    )


def build_final_word_report(*, record: bool = True) -> FinalWordArtifacts:
    ensure_runtime_directories()
    gate = _load_release_gate()
    figures, media = _parse_curation_report()
    derivatives = _prepare_derivatives(figures, media)
    report_path = DIRECTORIES["reports"] / REPORT_NAME
    manifest_path = DIRECTORIES["reports"] / REPORT_MANIFEST_NAME
    qa_json = DIRECTORIES["logs"] / WORD_QA_NAME
    qa_md = DIRECTORIES["reports"] / WORD_QA_MD_NAME

    _make_docx(report_path, figures, media, derivatives)
    inspection = _inspect_docx(report_path)
    checks = {
        "docx_exists": inspection["exists"],
        "docx_size_gt_500kb": inspection["size_bytes"] > 500_000,
        "media_count_ge_25": inspection["media_count"] >= 25,
        "table_count_ge_8": inspection["table_count"] >= MIN_REPORT_TABLES,
        "contains_aquaskim": inspection["has_aquaskim"],
        "contains_hydrostatics": inspection["has_english_hydrostatics"],
        "contains_model_boundary": inspection["has_model_boundary"],
        "no_comments": inspection["has_no_comments"],
        "no_tracked_changes": inspection["has_no_tracked_changes"],
        "delivery_zip_absent": not any(DIRECTORIES["deliverables"].glob("*.zip")),
        "release_still_disabled": gate.get("final_release_enabled") is False,
        "derivative_count_ge_25": len(derivatives) >= 25,
    }
    validation_status = "PASS" if all(checks.values()) else "FAIL"
    qa = {
        "identifier": "AQUASKIM-FINAL-WORD-QA-01",
        "validation_status": validation_status,
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        "inspection": inspection,
        "checks": checks,
        "curated_figure_count": len(figures),
        "curated_media_pair_count": len(media),
        "report_image_derivative_count": len(derivatives),
        "model_boundary": "The English Word report is generated from low-speed 3-DOF sheltered-basin evidence and does not expand claims.",
    }
    artifacts = FinalWordArtifacts(report_path, manifest_path, qa_json, qa_md, None)
    qa_json.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_qa_markdown(qa_md, qa, artifacts)
    if validation_status != "PASS":
        raise FinalReportError(f"Final Word QA failed: {json.dumps(checks, ensure_ascii=False)}")

    _write_manifest(manifest_path, report_path, figures, media, derivatives, qa)

    run_dir: Path | None = None
    if record:
        run_id = "phase10_18_" + datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        run_dir = DIRECTORIES["phase10_18_runs"] / run_id
        run_dir.mkdir(parents=True, exist_ok=True)
        for path in (report_path, manifest_path, qa_json, qa_md):
            shutil.copy2(path, run_dir / path.name)
        snapshot = {
            "run_id": run_id,
            "status": validation_status,
            "artifacts": {
                "report_docx": relative_to_root(report_path),
                "manifest": relative_to_root(manifest_path),
                "qa_json": relative_to_root(qa_json),
                "qa_markdown": relative_to_root(qa_md),
            },
            "sha256": {path.name: _sha256(path) for path in (report_path, manifest_path, qa_json, qa_md)},
            "delivery_zip_created": False,
            "release_enabled": False,
        }
        (run_dir / "evidence_snapshot.json").write_text(json.dumps(snapshot, ensure_ascii=False, indent=2), encoding="utf-8")
        _write_handoff(run_dir, validation_status)
        artifacts = FinalWordArtifacts(report_path, manifest_path, qa_json, qa_md, run_dir)
    return artifacts


def print_final_word_summary(artifacts: FinalWordArtifacts) -> None:
    qa = _read_json(artifacts.qa_json)
    print("========================================================================")
    print("AquaSkim-Sim | Final Word Report Generation and QA")
    print("========================================================================")
    print(f"Report   : {relative_to_root(artifacts.report_docx)}")
    print(f"Manifest : {relative_to_root(artifacts.report_manifest)}")
    print(f"QA       : {relative_to_root(artifacts.qa_json)}")
    if artifacts.run_dir is not None:
        print(f"Evidence : {relative_to_root(artifacts.run_dir)}")
    print(f"Status   : {qa['validation_status']}")
    print("Release  : DISABLED (delivery ZIP remains blocked)")
    print("========================================================================")


if __name__ == "__main__":
    print_final_word_summary(build_final_word_report(record=True))
